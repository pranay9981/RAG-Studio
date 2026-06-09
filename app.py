import streamlit as st

st.set_page_config(page_title="Top 6 RAG Architectures", page_icon="🤖", layout="wide")

import tempfile
import os
import base64
import time
import streamlit.components.v1 as components
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from core.shared_services import services
from architectures.hybrid_rag import HybridRAGPipeline
from architectures.graph_rag import GraphRAGPipeline
from architectures.agentic_rag import AgenticRAGPipeline
from architectures.corrective_rag import CorrectiveRAGPipeline
from architectures.multimodal_rag import MultimodalRAGPipeline
from architectures.multilingual_rag import MultilingualRAGPipeline

# ── Architecture metadata ────────────────────────────────────────────────────

ARCH_INFO = {
    "01 Hybrid RAG (Dense + Sparse)": {
        "icon": "🔀",
        "tagline": "Dense vector search + BM25 keyword search fused via Reciprocal Rank Fusion",
        "how": (
            "Runs **two retrievers in parallel**: ChromaDB (semantic) and BM25 (keyword). "
            "Their ranked lists are merged with RRF — documents that rank highly in *both* "
            "methods get a boosted score and rise to the top."
        ),
        "best_for": "General-purpose documents — best accuracy across mixed query types",
        "state_key": "hybrid_pipeline",
        "label": "Hybrid RAG (Dense + Sparse Fusion)",
    },
    "02 Graph RAG (Knowledge Graphs)": {
        "icon": "🕸️",
        "tagline": "LLM-extracted entity/relationship graph + vector fallback",
        "how": (
            "Uses Gemini to extract **(entity → relationship → entity)** triples from every chunk "
            "and builds a NetworkX knowledge graph. At query time it walks the graph for matching "
            "entities, collects relationship context, then combines it with dense vector results."
        ),
        "best_for": "Documents rich in named entities and relationships (research papers, reports)",
        "state_key": "graph_pipeline",
        "label": "Graph RAG (Entities + Knowledge Graph)",
    },
    "03 Agentic RAG (LangGraph)": {
        "icon": "🤖",
        "tagline": "LangGraph planner routes queries to vector search, web search, or direct answer",
        "how": (
            "A **3-node LangGraph state machine**: Planner → Tool Executor → Reasoner. "
            "The Planner decides whether to use `VECTOR_SEARCH` (internal docs), `WEB_SEARCH` "
            "(DuckDuckGo), or answer directly. The Reasoner synthesises the final answer."
        ),
        "best_for": "Queries that may need web context or multi-step reasoning",
        "state_key": "agentic_pipeline",
        "label": "Agentic RAG (Planner → Tools → Reasoner)",
    },
    "04 Corrective RAG (CRAG)": {
        "icon": "✅",
        "tagline": "Evaluator grades retrieved docs; rewrites query and falls back to web if needed",
        "how": (
            "A **5-node LangGraph workflow**: Retrieve → Evaluate → Route → Generate. "
            "The Evaluator grades retrieved docs as `CORRECT`, `AMBIGUOUS`, or `INCORRECT`. "
            "`AMBIGUOUS` triggers query rewriting; `INCORRECT` triggers a full web search fallback."
        ),
        "best_for": "When retrieval quality is uncertain or documents may not cover the query",
        "state_key": "crag_pipeline",
        "label": "Corrective RAG (Retrieve → Evaluate → Correct)",
    },
    "05 Multimodal RAG (Vision + Text)": {
        "icon": "🖼️",
        "tagline": "Stores images in metadata; sends text + image to Gemini vision for answers",
        "how": (
            "When an image is uploaded, Gemini generates a text summary for embedding. "
            "The original base64 image is stored in ChromaDB metadata. At query time, retrieved "
            "chunks that have an attached image include both text and the raw image in the "
            "multimodal Gemini message."
        ),
        "best_for": "Documents with figures, charts, screenshots, or mixed image/text content",
        "state_key": "multimodal_pipeline",
        "label": "Multimodal RAG (Vision + Text)",
    },
    "06 Multilingual RAG (BGE-M3)": {
        "icon": "🌍",
        "tagline": "Cross-lingual embedding space — query in any language, retrieve from any language",
        "how": (
            "Uses a multilingual sentence-transformer model so all languages share the same "
            "vector space. Retrieval is language-agnostic and the generation prompt instructs "
            "Gemini to answer in the **same language as the query**."
        ),
        "best_for": "Multilingual documents or when users may query in different languages",
        "state_key": "multilingual_pipeline",
        "label": "Multilingual RAG (Cross-lingual Retrieval)",
    },
}

ARCH_KEYS = list(ARCH_INFO.keys())

# ── Session state initialisation ─────────────────────────────────────────────

if "hybrid_pipeline"      not in st.session_state:
    st.session_state.hybrid_pipeline      = HybridRAGPipeline()
if "graph_pipeline"       not in st.session_state:
    st.session_state.graph_pipeline       = GraphRAGPipeline()
if "agentic_pipeline"     not in st.session_state:
    st.session_state.agentic_pipeline     = AgenticRAGPipeline()
if "crag_pipeline"        not in st.session_state:
    st.session_state.crag_pipeline        = CorrectiveRAGPipeline()
if "multimodal_pipeline"  not in st.session_state:
    st.session_state.multimodal_pipeline  = MultimodalRAGPipeline()
if "multilingual_pipeline" not in st.session_state:
    st.session_state.multilingual_pipeline = MultilingualRAGPipeline()
if "messages"             not in st.session_state:
    st.session_state.messages             = []
if "ingested_archs"       not in st.session_state:
    st.session_state.ingested_archs       = set()

# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.title("⚙️ Controls")

    selected_arch = st.selectbox("RAG Architecture", ARCH_KEYS)
    compare_mode  = st.checkbox("🔍 Compare All 6 Architectures")

    st.divider()

    # Document status
    st.markdown("**Document Status**")
    for key in ARCH_KEYS:
        info   = ARCH_INFO[key]
        loaded = key in st.session_state.ingested_archs
        dot    = "🟢" if loaded else "⚪"
        short  = key[:18]
        st.markdown(f"{dot} {short}")

    st.divider()

    # File upload + ingest
    uploaded_file = st.file_uploader(
        "Upload Document",
        type=["pdf", "txt", "docx", "png", "jpg", "jpeg"],
        help="PDF, TXT, DOCX, or image (PNG/JPG)",
    )

    if st.button("⬆️ Ingest Document", use_container_width=True, disabled=uploaded_file is None):
        with st.spinner("Processing…"):
            try:
                name = uploaded_file.name.lower()
                if name.endswith(".pdf"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
                        f.write(uploaded_file.getvalue())
                        tmp = f.name
                    docs = services.load_pdf(tmp)
                    os.remove(tmp)

                elif name.endswith(".txt"):
                    text  = uploaded_file.getvalue().decode("utf-8", errors="replace")
                    chunks = services.text_splitter.split_text(text)
                    docs  = [Document(page_content=c, metadata={"source": uploaded_file.name, "type": "txt"}) for c in chunks]

                elif name.endswith(".docx"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                        f.write(uploaded_file.getvalue())
                        tmp = f.name
                    from docx import Document as DocxDocument
                    docx   = DocxDocument(tmp)
                    text   = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
                    os.remove(tmp)
                    chunks = services.text_splitter.split_text(text)
                    docs   = [Document(page_content=c, metadata={"source": uploaded_file.name, "type": "docx"}) for c in chunks]

                else:
                    b64  = base64.b64encode(uploaded_file.getvalue()).decode("utf-8")
                    msg  = HumanMessage(content=[
                        {"type": "text", "text": "Describe this image in detail."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                    ])
                    summary = services.extract_response_text(services.llm.invoke([msg]))
                    docs = [Document(
                        page_content=f"Image Description: {summary}",
                        metadata={"source": uploaded_file.name, "type": "image", "image_base64": b64},
                    )]

                if not docs:
                    st.error("No content could be extracted.")
                else:
                    targets = ARCH_KEYS if compare_mode else [selected_arch]
                    for arch_key in targets:
                        sk = ARCH_INFO[arch_key]["state_key"]
                        st.session_state[sk].ingest(docs)
                        st.session_state.ingested_archs.add(arch_key)

                    scope = "all 6 architectures" if compare_mode else ARCH_INFO[selected_arch]["label"]
                    st.success(f"✅ {len(docs)} chunks → {scope}")

            except Exception as e:
                st.error(f"Ingestion failed: {e}")

    st.divider()

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ Clear Chat", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    with col2:
        if st.button("🔄 Reset Docs", use_container_width=True):
            for key in ARCH_KEYS:
                sk = ARCH_INFO[key]["state_key"]
                st.session_state[sk].reset()
            st.session_state.ingested_archs = set()
            st.success("All document collections cleared.")
            st.rerun()

# ── Main area ─────────────────────────────────────────────────────────────────

st.title("Top 6 RAG Architectures in 2026")

# Architecture info card
info = ARCH_INFO[selected_arch]
with st.expander(f"{info['icon']} **{selected_arch}** — {info['tagline']}", expanded=False):
    col_a, col_b = st.columns([3, 1])
    with col_a:
        st.markdown("**How it works**")
        st.markdown(info["how"])
    with col_b:
        st.markdown("**Best for**")
        st.info(info["best_for"])

# Knowledge graph visualization (Graph RAG only, after ingest)
if selected_arch == "02 Graph RAG (Knowledge Graphs)" and selected_arch in st.session_state.ingested_archs:
    graph_pipeline = st.session_state.graph_pipeline
    node_count = graph_pipeline.graph.number_of_nodes()
    edge_count = graph_pipeline.graph.number_of_edges()
    with st.expander(f"🕸️ Knowledge Graph — {node_count} entities · {edge_count} relationships", expanded=True):
        if node_count == 0:
            st.info("No entities extracted yet. Try ingesting a document with rich named entities.")
        else:
            html = graph_pipeline.render_graph_html()
            if html:
                components.html(html, height=540, scrolling=False)
                st.caption("Drag nodes to rearrange · Scroll to zoom · Hover for relationship labels")

st.divider()

# Chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Chat input
if prompt := st.chat_input("Ask a question about your document…"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user", avatar="🧑"):
        st.markdown(prompt)

    with st.chat_message("assistant", avatar="🤖"):
        if compare_mode:
            st.markdown("### 🔍 Comparing All 6 Architectures")

            timings = {}
            answers = {}

            for arch_key in ARCH_KEYS:
                icon_a = ARCH_INFO[arch_key]["icon"]
                sk     = ARCH_INFO[arch_key]["state_key"]
                with st.status(f"{icon_a} {arch_key}…", expanded=False) as s:
                    t0 = time.time()
                    try:
                        answers[arch_key] = st.session_state[sk].query(prompt)
                    except Exception as e:
                        answers[arch_key] = f"⚠️ Error: {e}"
                    elapsed_i = time.time() - t0
                    timings[arch_key] = elapsed_i
                    s.update(label=f"{icon_a} {arch_key} — {elapsed_i:.2f}s", state="complete")

            # Timing summary table
            timing_rows = "| Architecture | Time |\n|---|---|\n"
            for k in ARCH_KEYS:
                timing_rows += f"| {ARCH_INFO[k]['icon']} {k} | `{timings[k]:.2f}s` |\n"
            st.markdown(timing_rows)

            st.divider()

            # Individual results
            for arch_key in ARCH_KEYS:
                icon_a = ARCH_INFO[arch_key]["icon"]
                with st.expander(f"{icon_a} **{arch_key}** — `{timings[arch_key]:.2f}s`", expanded=False):
                    st.markdown(answers[arch_key])

            response = "Comparison complete — expand each tab above to read the answers."

        else:
            collected_tokens = []
            answer = ""
            elapsed = 0.0

            with st.status(f"🧠 {info['icon']} {selected_arch}…", expanded=True) as status:
                def on_event(event):
                    kind, content = event
                    if kind == "step":
                        st.write(f"✅ {content}")
                    elif kind == "token":
                        collected_tokens.append(content)

                t0 = time.time()
                try:
                    sk       = info["state_key"]
                    pipeline = st.session_state[sk]
                    answer   = pipeline.query(prompt, on_step=on_event)
                    elapsed  = time.time() - t0
                    status.update(label=f"✅ Done in {elapsed:.2f}s", state="complete", expanded=False)
                except Exception as e:
                    elapsed = time.time() - t0
                    answer  = f"⚠️ Pipeline error: {e}"
                    status.update(label="❌ Error", state="error", expanded=False)
                    st.error(answer)

            final_text = "".join(collected_tokens) if collected_tokens else answer
            st.caption(f"{info['icon']} **{info['label']}** · `{elapsed:.2f}s`")
            st.markdown(final_text)
            response = final_text

    st.session_state.messages.append({"role": "assistant", "content": response})
