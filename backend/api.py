import os
import re
import json
import time
import base64
import asyncio
import socket
import tempfile
import threading
import urllib.request
from queue import Queue, Empty
from typing import Optional, List

from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from backend.session_manager import session, ARCH_KEYS, STATE_KEY_MAP, ARCH_INFO
from core.shared_services import services
from core.adaptive_db import adaptive_db
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage

# ── App ───────────────────────────────────────────────────────────────────────

app = FastAPI(title="RAG System API", version="4.0.0", docs_url="/docs")

_raw_origins = os.environ.get(
    "FRONTEND_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000"
)
_ALLOWED_ORIGINS = [o.strip() for o in _raw_origins.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Pydantic models ───────────────────────────────────────────────────────────

class CompareRequest(BaseModel):
    query: str
    session_id: str = "default"

class EvalRequest(BaseModel):
    query: str
    answer: str
    arch_key: str = ""
    sources: List[dict] = []

class FeedbackRequest(BaseModel):
    query: str
    arch_key: str
    chunk_ids: List[str] = []
    rating: int  # 1 = thumbs up, -1 = thumbs down
    message_id: str = ""

class DeleteDocumentRequest(BaseModel):
    source: str  # filename label — deletes from ALL architectures

class ApiKeyRequest(BaseModel):
    api_key: str

# ── Helpers ───────────────────────────────────────────────────────────────────

def fetch_url_text(url: str) -> str:
    from urllib.parse import urlparse
    import ipaddress
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(status_code=400, detail="Only http and https URLs are supported")
    host = parsed.hostname or ""
    if not host:
        raise HTTPException(status_code=400, detail="Could not parse hostname")
    # Resolve DNS and validate every returned IP
    try:
        addr_infos = socket.getaddrinfo(host, None)
        for *_, sockaddr in addr_infos:
            ip_str = sockaddr[0]
            try:
                addr = ipaddress.ip_address(ip_str)
                if addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved:
                    raise HTTPException(
                        status_code=400,
                        detail=f"URL resolves to a private or internal address ({ip_str})"
                    )
            except ValueError:
                pass
    except socket.gaierror as e:
        raise HTTPException(status_code=400, detail=f"DNS resolution failed: {e}")
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=15) as resp:
        html = resp.read().decode("utf-8", errors="replace")
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<style[^>]*>.*?</style>",   "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


async def process_file(file: UploadFile) -> List[Document]:
    content = await file.read()
    name = (file.filename or "").lower()
    source = file.filename or "upload"

    MAX_BYTES = 50 * 1024 * 1024  # 50 MB
    if len(content) > MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 50 MB limit")

    if name.endswith(".pdf"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
            f.write(content)
            tmp = f.name
        docs = services.load_pdf(tmp)
        os.remove(tmp)
        for doc in docs:
            doc.metadata["source"] = source  # override temp path with original filename
        return docs

    if name.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
        chunks = services.text_splitter.split_text(text)
        return [Document(page_content=c, metadata={"source": source, "type": "txt"}) for c in chunks]

    if name.endswith(".docx"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
            f.write(content)
            tmp = f.name
        from docx import Document as DocxDocument
        docx = DocxDocument(tmp)
        text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
        os.remove(tmp)
        chunks = services.text_splitter.split_text(text)
        return [Document(page_content=c, metadata={"source": source, "type": "docx"}) for c in chunks]

    if name.endswith(".csv"):
        import pandas as pd
        import io as _io
        df = pd.read_csv(_io.BytesIO(content))
        schema = (
            f"Columns: {list(df.columns)}\n"
            f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
            f"Dtypes:\n" + "\n".join(f"  {col}: {dtype}" for col, dtype in df.dtypes.items())
        )
        preview = df.head(500).to_csv(index=False)
        text = f"TABLE SCHEMA:\n{schema}\n\nDATA:\n{preview}"
        return [Document(
            page_content=text,
            metadata={
                "source": source, "type": "csv",
                "columns": json.dumps(list(df.columns)),
                "rows": str(df.shape[0]),
            },
        )]

    if name.endswith((".xlsx", ".xls")):
        import pandas as pd
        import io as _io
        df = pd.read_excel(_io.BytesIO(content))
        schema = (
            f"Columns: {list(df.columns)}\n"
            f"Shape: {df.shape[0]} rows × {df.shape[1]} columns"
        )
        preview = df.head(500).to_csv(index=False)
        text = f"TABLE SCHEMA:\n{schema}\n\nDATA:\n{preview}"
        return [Document(
            page_content=text,
            metadata={
                "source": source, "type": "excel",
                "columns": json.dumps(list(df.columns)),
                "rows": str(df.shape[0]),
            },
        )]

    SUPPORTED_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".webp")
    if not name.endswith(SUPPORTED_IMAGE_EXTS):
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type. Supported: PDF, TXT, DOCX, CSV, XLSX, PNG, JPG, JPEG, GIF, WEBP"
        )

    b64 = base64.b64encode(content).decode("utf-8")
    msg = HumanMessage(content=[
        {"type": "text", "text": "Describe this image in detail."},
        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
    ])
    summary = services.extract_response_text(services.llm.invoke([msg]))
    import uuid as _uuid
    img_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'adaptive_data', 'images')
    os.makedirs(img_dir, exist_ok=True)
    img_filename = f"{_uuid.uuid4().hex}.b64"
    img_path = os.path.join(img_dir, img_filename)
    with open(img_path, 'w') as fh:
        fh.write(b64)
    return [Document(
        page_content=f"Image Description: {summary}",
        metadata={"source": source, "type": "image", "image_path": img_path},
    )]

# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "architectures": len(ARCH_KEYS)}


@app.get("/api/architectures")
async def get_architectures():
    return list(ARCH_INFO.values())


@app.get("/api/sessions/{session_id}")
async def get_session(session_id: str):
    return {
        "session_id": session_id,
        "ingested_archs": list(session.ingested_archs),
        "doc_library": session.doc_library,
        "history_count": len(session.history),
    }


@app.post("/api/ingest")
async def ingest_document(
    session_id: str = Form(default="default"),
    arch_keys: str = Form(...),
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None),
):
    if arch_keys.strip() == "all":
        target_keys = ARCH_KEYS
    else:
        try:
            parsed = json.loads(arch_keys)
            target_keys = parsed if isinstance(parsed, list) else [parsed]
        except Exception:
            target_keys = [arch_keys.strip()]

    docs: List[Document] = []
    source_name = ""

    if file is not None and file.filename:
        source_name = file.filename
        docs = await process_file(file)
    elif url and url.strip():
        source_name = url.strip()
        try:
            raw = fetch_url_text(source_name)
            if len(raw) < 100:
                raise HTTPException(status_code=422, detail="Could not extract meaningful text from URL")
            chunks = services.text_splitter.split_text(raw)
            docs = [Document(page_content=c, metadata={"source": source_name, "type": "url"}) for c in chunks]
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"URL fetch failed: {e}")
    else:
        raise HTTPException(status_code=400, detail="Provide either a file or a URL")

    if not docs:
        raise HTTPException(status_code=422, detail="No content could be extracted")

    # Reject re-upload of already-ingested source
    existing_sources = {d["name"] for d in session.doc_library}
    if source_name in existing_sources:
        raise HTTPException(
            status_code=409,
            detail=f"'{source_name}' is already ingested. Delete it first to re-upload."
        )

    # Parent-child chunking: small children for retrieval, parent text stored in metadata for generation
    docs = services.create_parent_child_documents(docs)

    ingested: List[str] = []
    for arch_key in target_keys:
        state_key = STATE_KEY_MAP.get(arch_key)
        if not state_key:
            continue
        pipeline = session.get_pipeline(state_key)
        if pipeline:
            try:
                pipeline.ingest(docs)
                session.add_ingested_arch(arch_key)
                ingested.append(arch_key)
            except Exception as exc:
                print(f"[ingest] {arch_key} failed: {exc}")

    if ingested:
        session.append_doc({"name": source_name, "chunks": len(docs)})
    return {"chunks": len(docs), "source": source_name, "architectures": ingested}


@app.get("/api/query")
async def query_stream(
    session_id: str = Query(default="default"),
    query: str = Query(...),
    arch_key: str = Query(...),
):
    state_key = STATE_KEY_MAP.get(arch_key)
    if not state_key:
        raise HTTPException(status_code=400, detail=f"Unknown architecture: {arch_key}")

    pipeline = session.get_pipeline(state_key)
    if not pipeline:
        raise HTTPException(status_code=400, detail="Pipeline not found")

    # Check semantic cache
    try:
        query_embedding = services.embeddings.embed_query(query)
        cached = adaptive_db.find_similar_query(query_embedding, arch_key)
    except Exception:
        cached = None
        query_embedding = []  # empty list, not None — prevents None-type errors downstream

    if cached:
        async def cached_generator():
            sim = cached["similarity"]
            yield f'data: {json.dumps({"type": "step", "content": f"⚡ Semantic cache hit (similarity {sim}) — reusing previous answer"})}\n\n'
            if cached["sources"]:
                yield f'data: {json.dumps({"type": "sources", "content": cached["sources"]})}\n\n'
            # Stream cached answer token by token for visual continuity
            for word in cached["answer"].split(" "):
                yield f'data: {json.dumps({"type": "token", "content": word + " "})}\n\n'
            yield f'data: {json.dumps({"type": "done", "answer": cached["answer"], "elapsed": 0.001, "cached": True})}\n\n'
            adaptive_db.store_query_analytics(arch_key, query, 0.001, cached=True)

        return StreamingResponse(
            cached_generator(),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
        )

    q: "Queue[tuple]" = Queue()
    t0 = time.time()
    collected_sources: List[dict] = []

    def on_step(event):
        q.put(event)

    def run_pipeline():
        try:
            result = pipeline.query(query, on_step=on_step)
            elapsed = time.time() - t0
            q.put(("done", {"answer": result, "elapsed": round(elapsed, 3)}))
        except Exception as e:
            q.put(("error", {"message": str(e)}))

    thread = threading.Thread(target=run_pipeline, daemon=True)
    thread.start()

    async def event_generator():
        collected_answer = ""
        elapsed = 0.0

        while True:
            try:
                event = q.get_nowait()
                kind, content = event

                if kind == "sources":
                    collected_sources.extend(content)
                    payload = json.dumps({"type": "sources", "content": content})
                elif kind == "token":
                    payload = json.dumps({"type": "token", "content": content})
                elif kind == "step":
                    payload = json.dumps({"type": "step", "content": content})
                elif kind == "done":
                    collected_answer = content.get("answer", "")
                    elapsed = content.get("elapsed", 0)
                    payload = json.dumps({
                        "type": "done",
                        "answer": collected_answer,
                        "elapsed": elapsed,
                        "cached": False,
                    })
                elif kind == "error":
                    payload = json.dumps({"type": "error", "content": content.get("message", "Unknown error")})
                else:
                    payload = json.dumps({"type": kind, "content": str(content)})

                yield f"data: {payload}\n\n"

                if kind in ("done", "error"):
                    if kind == "done":
                        session.append_history({
                            "query":   query,
                            "arch":    arch_key,
                            "elapsed": elapsed,
                            "answer":  collected_answer[:120],
                        })
                        adaptive_db.store_query_analytics(arch_key, query, elapsed)
                        # Store in semantic cache for future similar queries
                        if collected_answer and query_embedding:
                            try:
                                adaptive_db.store_query_cache(
                                    arch_key, query, query_embedding,
                                    collected_answer, collected_sources[:5],
                                )
                            except Exception as e:
                                print(f"[cache] store_query_cache failed: {e}")
                    break

            except Empty:
                if not thread.is_alive():
                    # Drain residual events — thread is done, no more items will be added
                    while True:
                        try:
                            event = q.get_nowait()
                        except Empty:
                            return
                        kind, content = event
                        if kind == "sources":
                            collected_sources.extend(content)
                            payload = json.dumps({"type": "sources", "content": content})
                        elif kind == "token":
                            payload = json.dumps({"type": "token", "content": content})
                        elif kind == "step":
                            payload = json.dumps({"type": "step", "content": content})
                        elif kind == "done":
                            collected_answer = content.get("answer", "")
                            elapsed = content.get("elapsed", 0)
                            payload = json.dumps({"type": "done", "answer": collected_answer, "elapsed": elapsed, "cached": False})
                        elif kind == "error":
                            payload = json.dumps({"type": "error", "content": content.get("message", "Unknown error")})
                        else:
                            payload = json.dumps({"type": kind, "content": str(content)})
                        yield f"data: {payload}\n\n"
                        if kind in ("done", "error"):
                            if kind == "done":
                                session.append_history({"query": query, "arch": arch_key, "elapsed": elapsed, "answer": collected_answer[:120]})
                                adaptive_db.store_query_analytics(arch_key, query, elapsed)
                                if collected_answer and query_embedding:
                                    try:
                                        adaptive_db.store_query_cache(arch_key, query, query_embedding, collected_answer, collected_sources[:5])
                                    except Exception as e:
                                        print(f"[cache] store_query_cache failed: {e}")
                            return
                    return
                await asyncio.sleep(0.01)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":     "no-cache",
            "Connection":        "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/compare")
async def compare(request: CompareRequest):
    result_list = [None] * len(ARCH_KEYS)

    def run_one(i: int, arch_key: str, state_key: str):
        pipeline = session.get_pipeline(state_key)
        start = time.time()
        try:
            answer = pipeline.query(request.query)
            result_list[i] = {
                "arch_key": arch_key,
                "answer": answer,
                "elapsed": round(time.time() - start, 3),
            }
        except Exception as e:
            result_list[i] = {
                "arch_key": arch_key,
                "answer": "",
                "elapsed": round(time.time() - start, 3),
                "error": str(e),
            }

    threads = []
    for i, arch_key in enumerate(ARCH_KEYS):
        t = threading.Thread(
            target=run_one, args=(i, arch_key, STATE_KEY_MAP[arch_key]), daemon=True
        )
        threads.append(t)
        t.start()

    for t in threads:
        t.join(timeout=120)

    return {"results": [r for r in result_list if r is not None]}


@app.post("/api/evaluate")
async def evaluate_answer(request: EvalRequest):
    context = "\n".join(s.get("text", "") for s in request.sources) if request.sources else ""
    scores = {"faithfulness": 0, "relevance": 0, "context_precision": 0, "context_recall": 0}
    eval_ok = False

    # ── Step 1: RAGAS-style faithfulness — extract claims then verify each ────
    faithfulness_score = 0
    try:
        claims_prompt = f"""Break this answer into individual factual claims. Each claim = one verifiable statement.
Answer: {request.answer[:600]}
Output ONLY a JSON array of strings: ["claim1", "claim2", ...]"""
        claims_resp = services.llm.invoke(claims_prompt)
        claims_text = services.extract_response_text(claims_resp)
        cm = re.search(r"\[.*?\]", claims_text, re.DOTALL)
        claims = json.loads(cm.group()) if cm else []

        if claims and context:
            verify_prompt = f"""For each claim, output true if it is SUPPORTED by the context, false otherwise.
Context: {context[:1500]}
Claims: {json.dumps(claims[:10])}
Output ONLY a JSON array of booleans (one per claim): [true, false, ...]"""
            verify_resp = services.llm.invoke(verify_prompt)
            verify_text = services.extract_response_text(verify_resp)
            vm = re.search(r"\[.*?\]", verify_text, re.DOTALL)
            supported = json.loads(vm.group()) if vm else []
            if supported and claims:
                faithfulness_score = round(sum(1 for s in supported if s) / len(claims) * 10)
    except Exception as e:
        print(f"[evaluate] faithfulness step failed: {e}")

    # ── Step 2: relevance, context_precision, context_recall in one call ─────
    try:
        other_prompt = f"""You are an expert RAG evaluator. Score on three dimensions.

Question: {request.query}
Retrieved Context: {context[:1200] if context else "N/A"}
Generated Answer: {request.answer[:600]}

Score each 0-10:
1. Relevance: Does the answer directly address the question? (0=off-topic, 10=perfect)
2. Context Precision: Were the retrieved chunks relevant to the query? (0=all irrelevant, 10=all relevant)
3. Context Recall: Did the context contain all information needed to fully answer? (0=missing key info, 10=complete)

Output ONLY valid JSON: {{"relevance": X, "context_precision": X, "context_recall": X}}"""

        resp = services.llm.invoke(other_prompt)
        text = services.extract_response_text(resp)
        m = re.search(r"\{[^}]+\}", text)
        if m:
            raw = json.loads(m.group())
            scores = {
                "faithfulness": faithfulness_score,
                "relevance": max(0, min(10, int(raw.get("relevance", 0)))),
                "context_precision": max(0, min(10, int(raw.get("context_precision", 0)))),
                "context_recall": max(0, min(10, int(raw.get("context_recall", 0)))),
            }
            eval_ok = True
    except Exception as e:
        print(f"[evaluate] other metrics step failed: {e}")
        if faithfulness_score:
            scores["faithfulness"] = faithfulness_score
            eval_ok = True

    # Only store analytics when at least one LLM step succeeded
    if eval_ok and request.arch_key:
        try:
            adaptive_db.store_eval_analytics(request.arch_key, request.query, scores)
        except Exception as e:
            print(f"[evaluate] store_eval_analytics failed: {e}")

    return scores


@app.post("/api/feedback")
async def submit_feedback(request: FeedbackRequest):
    if request.rating not in (1, -1):
        raise HTTPException(status_code=400, detail="Rating must be 1 (up) or -1 (down)")
    adaptive_db.store_feedback(
        request.query, request.arch_key, request.chunk_ids, request.rating
    )
    return {"status": "ok"}


@app.get("/api/analytics")
async def get_analytics():
    return {
        "data": adaptive_db.get_analytics(),
        "recent": adaptive_db.get_recent_queries(20),
    }


@app.get("/api/graph")
async def get_graph():
    pipeline = session.get_pipeline("graph_pipeline")
    if not pipeline:
        return {"html": ""}
    try:
        html = pipeline.render_graph_html()
        return {"html": html or ""}
    except Exception as e:
        return {"html": "", "error": str(e)}


@app.get("/api/history")
async def get_history(session_id: str = Query(default="default")):
    return {"history": session.history[-20:]}


@app.get("/api/documents")
async def list_documents(arch_key: str = Query(...)):
    state_key = STATE_KEY_MAP.get(arch_key)
    if not state_key:
        raise HTTPException(status_code=400, detail="Unknown architecture")
    pipeline = session.get_pipeline(state_key)
    if not pipeline or not hasattr(pipeline, "collection"):
        raise HTTPException(status_code=400, detail="Pipeline not found")
    try:
        result = pipeline.collection.get(include=["metadatas"])
        sources: dict = {}
        for meta in (result["metadatas"] or []):
            raw = (meta or {}).get("source", "Unknown")
            label = raw.split("/")[-1].split("\\")[-1]
            sources[label] = sources.get(label, 0) + 1
        return {"documents": [{"source": k, "chunks": v} for k, v in sources.items()]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/documents")
async def delete_document(request: DeleteDocumentRequest):
    total_deleted = 0
    for arch_key, state_key in STATE_KEY_MAP.items():
        pipeline = session.get_pipeline(state_key)
        if not pipeline or not hasattr(pipeline, "collection"):
            continue
        try:
            result = pipeline.collection.get(include=["metadatas", "ids"])
            ids_to_delete = [
                doc_id for doc_id, meta in zip(result["ids"] or [], result["metadatas"] or [])
                if (meta or {}).get("source", "Unknown").split("/")[-1].split("\\")[-1] == request.source
            ]
            if ids_to_delete:
                pipeline.collection.delete(ids=ids_to_delete)
                total_deleted += len(ids_to_delete)
        except Exception as e:
            print(f"[delete_document] {arch_key} failed: {e}")

    # Remove from structured RAG table store
    structured = session.get_pipeline("structured_pipeline")
    if structured and hasattr(structured, "_table_store"):
        keys = [k for k in structured._table_store
                if k.split("/")[-1].split("\\")[-1] == request.source]
        for k in keys:
            del structured._table_store[k]

    # Update session doc_library and ingested_archs
    session.filter_doc_library(
        lambda d: d.get("name", "").split("/")[-1].split("\\")[-1] != request.source
    )
    for arch_key, state_key in STATE_KEY_MAP.items():
        p = session.get_pipeline(state_key)
        if p and hasattr(p, "collection") and p.collection.count() == 0:
            session.ingested_archs.discard(arch_key)

    return {"deleted": total_deleted, "source": request.source}


@app.get("/api/config/status")
async def get_config_status():
    return {"has_key": bool(os.environ.get("GROQ_API_KEY", "").strip())}


@app.post("/api/config/apikey")
async def set_api_key(request: ApiKeyRequest):
    if not request.api_key.strip():
        raise HTTPException(status_code=400, detail="API key cannot be empty")
    key = request.api_key.strip()
    os.environ["GROQ_API_KEY"] = key

    try:
        from langchain_groq import ChatGroq
        services.llm = ChatGroq(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.2,
            max_tokens=1024,
        )
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to initialize LLM: {e}")


@app.delete("/api/sessions/{session_id}")
async def reset_session(session_id: str):
    session.reset()
    return {"status": "reset", "session_id": session_id}
