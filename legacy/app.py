import streamlit as st

st.set_page_config(page_title="Top 8 RAG Architectures", page_icon="🤖", layout="wide")

import tempfile
import os
import base64
import time
import re
import json
import urllib.request
import streamlit.components.v1 as components
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from core.shared_services import services
from architectures.hybrid_rag import HybridRAGPipeline
from architectures.graph_rag import GraphRAGPipeline
from architectures.agentic_rag import AgenticRAGPipeline
from architectures.corrective_rag import CorrectiveRAGPipeline
from architectures.multimodal_rag import MultimodalRAGPipeline
from architectures.multilingual_rag import MultilingualRAGPipeline
from architectures.rag_fusion import RAGFusionPipeline
from architectures.hyde_rag import HyDERAGPipeline

# ── Architecture metadata ────────────────────────────────────────────────────

ARCH_INFO = {
    "01 Hybrid RAG (Dense + Sparse)": {
        "icon": "🔀",
        "tagline": "Dense vector search + BM25 keyword search fused via Reciprocal Rank Fusion",
        "how": (
            "Runs **two retrievers in parallel**: ChromaDB (semantic) and BM25 (keyword). "
            "Their ranked lists are merged with RRF — documents that rank highly in *both* "
            "methods get a boosted score and rise to the top. A cross-encoder then re-ranks "
            "the fused results for maximum precision."
        ),
        "best_for": "General-purpose documents — best accuracy across mixed query types",
        "state_key": "hybrid_pipeline",
        "label": "Hybrid RAG (Dense + Sparse + Re-ranking)",
    },
    "02 Graph RAG (Knowledge Graphs)": {
        "icon": "🕸️",
        "tagline": "LLM-extracted entity/relationship graph + vector fallback",
        "how": (
            "Uses Gemini to extract **(entity → relationship → entity)** triples from every chunk "
            "and builds a NetworkX knowledge graph. At query time it walks the graph for matching "
            "entities, collects relationship context, then combines it with dense vector results."
        ),
        "best_for": "Documents rich in named entities and relationships (research papers, reports)",
        "state_key": "graph_pipeline",
        "label": "Graph RAG (Entities + Knowledge Graph)",
    },
    "03 Agentic RAG (LangGraph)": {
        "icon": "🤖",
        "tagline": "LangGraph planner routes queries to vector search, web search, or direct answer",
        "how": (
            "A **3-node LangGraph state machine**: Planner → Tool Executor → Reasoner. "
            "The Planner decides whether to use `VECTOR_SEARCH` (internal docs), `WEB_SEARCH` "
            "(DuckDuckGo), or answer directly. The Reasoner synthesises the final answer."
        ),
        "best_for": "Queries that may need web context or multi-step reasoning",
        "state_key": "agentic_pipeline",
        "label": "Agentic RAG (Planner → Tools → Reasoner)",
    },
    "04 Corrective RAG (CRAG)": {
        "icon": "✅",
        "tagline": "Evaluator grades retrieved docs; rewrites query and falls back to web if needed",
        "how": (
            "A **5-node LangGraph workflow**: Retrieve → Evaluate → Route → Generate. "
            "The Evaluator grades retrieved docs as `CORRECT`, `AMBIGUOUS`, or `INCORRECT`. "
            "`AMBIGUOUS` triggers query rewriting; `INCORRECT` triggers a full web search fallback."
        ),
        "best_for": "When retrieval quality is uncertain or documents may not cover the query",
        "state_key": "crag_pipeline",
        "label": "Corrective RAG (Retrieve → Evaluate → Correct)",
    },
    "05 Multimodal RAG (Vision + Text)": {
        "icon": "🖼️",
        "tagline": "Stores images in metadata; sends text + image to Gemini vision for answers",
        "how": (
            "When an image is uploaded, Gemini generates a text summary for embedding. "
            "The original base64 image is stored in ChromaDB metadata. At query time, retrieved "
            "chunks that have an attached image include both text and the raw image in the "
            "multimodal Gemini message."
        ),
        "best_for": "Documents with figures, charts, screenshots, or mixed image/text content",
        "state_key": "multimodal_pipeline",
        "label": "Multimodal RAG (Vision + Text)",
    },
    "06 Multilingual RAG (BGE-M3)": {
        "icon": "🌍",
        "tagline": "Cross-lingual embedding space — query in any language, retrieve from any language",
        "how": (
            "Uses a multilingual sentence-transformer model so all languages share the same "
            "vector space. Retrieval is language-agnostic and a cross-encoder re-ranks results "
            "before the generation prompt instructs Gemini to answer in the **same language as the query**."
        ),
        "best_for": "Multilingual documents or when users may query in different languages",
        "state_key": "multilingual_pipeline",
        "label": "Multilingual RAG (Cross-lingual + Re-ranking)",
    },
    "07 RAG-Fusion (Query Expansion)": {
        "icon": "🔮",
        "tagline": "Expands your query into 4 sub-queries, retrieves for each, fuses with RRF",
        "how": (
            "Asks Gemini to generate **4 different phrasings** of your query, covering different "
            "angles. Each sub-query retrieves its own ranked list from ChromaDB. All four ranked "
            "lists are merged with **Reciprocal Rank Fusion** — documents that appear highly in "
            "multiple sub-query results get a compounding boost."
        ),
        "best_for": "Ambiguous or broad queries where a single phrasing might miss relevant docs",
        "state_key": "rag_fusion_pipeline",
        "label": "RAG-Fusion (Multi-Query + RRF)",
    },
    "08 HyDE RAG (Hypothetical Document)": {
        "icon": "💡",
        "tagline": "Generates a hypothetical answer first, uses it as the search query",
        "how": (
            "Instead of embedding your raw question, Gemini first generates a **hypothetical "
            "ideal answer** as if it were written in your document. That hypothetical is then "
            "embedded and used to retrieve the real chunks closest to it in vector space — "
            "bridging the vocabulary gap between short questions and long documents."
        ),
        "best_for": "Short or keyword-style queries where the question is worded very differently from the source text",
        "state_key": "hyde_pipeline",
        "label": "HyDE RAG (Hypothetical Document Embeddings)",
    },
}

ARCH_KEYS = list(ARCH_INFO.keys())

# ── Session state ─────────────────────────────────────────────────────────────

_defaults = {
    "hybrid_pipeline":       HybridRAGPipeline,
    "graph_pipeline":        GraphRAGPipeline,
    "agentic_pipeline":      AgenticRAGPipeline,
    "crag_pipeline":         CorrectiveRAGPipeline,
    "multimodal_pipeline":   MultimodalRAGPipeline,
    "multilingual_pipeline": MultilingualRAGPipeline,
    "rag_fusion_pipeline":   RAGFusionPipeline,
    "hyde_pipeline":         HyDERAGPipeline,
}
for key, cls in _defaults.items():
    if key not in st.session_state:
        st.session_state[key] = cls()

if "messages"       not in st.session_state:
    st.session_state.messages       = []
if "ingested_archs" not in st.session_state:
    st.session_state.ingested_archs = set()
if "doc_library"    not in st.session_state:
    st.session_state.doc_library    = []   # [{"name": str, "chunks": int}]
if "query_history"  not in st.session_state:
    st.session_state.query_history  = []   # [{"query", "arch", "elapsed", "answer"}]

# ── Utility functions ─────────────────────────────────────────────────────────

def fetch_url_text(url: str) -> str:
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=15) as resp:
        html = resp.read().decode("utf-8", errors="replace")
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<style[^>]*>.*?</style>",   "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def evaluate_answer(query: str, sources: list, answer: str) -> dict:
    context = "\n".join(s.get("text", "") for s in sources) if sources else ""
    prompt = f"""You are an expert RAG system evaluator. Score the answer on three dimensions.

Question: {query}
Retrieved Context (first 1500 chars): {context[:1500] if context else "N/A"}
Generated Answer: {answer[:800]}

Score each 0–10:
1. Faithfulness: Is the answer grounded in the context? (0=hallucinated, 10=fully supported)
2. Relevance: Does the answer directly address the question? (0=off-topic, 10=perfectly on-point)
3. Context Precision: Were the right chunks retrieved? (0=irrelevant, 10=perfectly targeted)

Output ONLY valid JSON: {{"faithfulness": X, "relevance": X, "context_precision": X}}"""
    try:
        response = services.llm.invoke(prompt)
        text = services.extract_response_text(response)
        m = re.search(r"\{[^}]+\}", text)
        if m:
            raw = json.loads(m.group())
            return {k: max(0, min(10, int(raw.get(k, 5)))) for k in ("faithfulness", "relevance", "context_precision")}
    except Exception:
        pass
    return {"faithfulness": 5, "relevance": 5, "context_precision": 5}


def export_chat_markdown() -> str:
    lines = ["# RAG Chat Export\n"]
    for msg in st.session_state.messages:
        role = "**You**" if msg["role"] == "user" else "**Assistant**"
        lines.append(f"{role}:\n{msg['content']}\n")
    return "\n---\n".join(lines)


def score_color(score: int) -> str:
    if score >= 7:
        return "🟢"
    if score >= 4:
        return "🟡"
    return "🔴"

# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.title("⚙️ Controls")

    selected_arch = st.selectbox("RAG Architecture", ARCH_KEYS)
    compare_mode  = st.checkbox("🔍 Compare All 8 Architectures")
    enable_eval   = st.checkbox("🧪 RAG Evaluation (extra LLM call)", value=False,
                                help="Scores each answer on Faithfulness, Relevance, and Context Precision using an LLM judge.")

    st.divider()

    # Document status indicators
    st.markdown("**Document Status**")
    cols_status = st.columns(2)
    for i, key in enumerate(ARCH_KEYS):
        info_item  = ARCH_INFO[key]
        loaded     = key in st.session_state.ingested_archs
        dot        = "🟢" if loaded else "⚪"
        short      = key[:16]
        cols_status[i % 2].markdown(f"{dot} {short}")

    st.divider()

    # Document library
    if st.session_state.doc_library:
        st.markdown("**📚 Ingested Documents**")
        seen_names = set()
        for doc in st.session_state.doc_library:
            if doc["name"] not in seen_names:
                seen_names.add(doc["name"])
                st.markdown(f"📄 `{doc['name']}` — {doc['chunks']} chunks")
        st.caption(f"{len(seen_names)} file(s) · {sum(d['chunks'] for d in st.session_state.doc_library)} total chunks")
        st.divider()

    # File upload
    uploaded_file = st.file_uploader(
        "Upload Document",
        type=["pdf", "txt", "docx", "png", "jpg", "jpeg"],
        help="PDF, TXT, DOCX, or image (PNG/JPG)",
    )

    # URL input
    url_input = st.text_input("Or enter a URL", placeholder="https://example.com/article")

    ingest_disabled = (uploaded_file is None) and (not url_input.strip())
    if st.button("⬆️ Ingest", use_container_width=True, disabled=ingest_disabled):
        with st.spinner("Processing…"):
            docs = []
            source_name = ""
            try:
                if uploaded_file is not None:
                    source_name = uploaded_file.name
                    name = uploaded_file.name.lower()

                    if name.endswith(".pdf"):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
                            f.write(uploaded_file.getvalue())
                            tmp = f.name
                        docs = services.load_pdf(tmp)
                        os.remove(tmp)

                    elif name.endswith(".txt"):
                        text   = uploaded_file.getvalue().decode("utf-8", errors="replace")
                        chunks = services.text_splitter.split_text(text)
                        docs   = [Document(page_content=c, metadata={"source": source_name, "type": "txt"}) for c in chunks]

                    elif name.endswith(".docx"):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                            f.write(uploaded_file.getvalue())
                            tmp = f.name
                        from docx import Document as DocxDocument
                        docx   = DocxDocument(tmp)
                        text   = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
                        os.remove(tmp)
                        chunks = services.text_splitter.split_text(text)
                        docs   = [Document(page_content=c, metadata={"source": source_name, "type": "docx"}) for c in chunks]

                    else:
                        b64     = base64.b64encode(uploaded_file.getvalue()).decode("utf-8")
                        msg_img = HumanMessage(content=[
                            {"type": "text", "text": "Describe this image in detail."},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                        ])
                        summary = services.extract_response_text(services.llm.invoke([msg_img]))
                        docs = [Document(
                            page_content=f"Image Description: {summary}",
                            metadata={"source": source_name, "type": "image", "image_base64": b64},
                        )]

                elif url_input.strip():
                    source_name = url_input.strip()
                    raw_text    = fetch_url_text(source_name)
                    if len(raw_text) < 100:
                        st.error("Could not extract meaningful text from that URL.")
                    else:
                        chunks = services.text_splitter.split_text(raw_text)
                        docs   = [Document(page_content=c, metadata={"source": source_name, "type": "url"}) for c in chunks]

                if not docs:
                    st.error("No content could be extracted.")
                else:
                    targets = ARCH_KEYS if compare_mode else [selected_arch]
                    for arch_key in targets:
                        sk = ARCH_INFO[arch_key]["state_key"]
                        st.session_state[sk].ingest(docs)
                        st.session_state.ingested_archs.add(arch_key)

                    # Track in doc library
                    st.session_state.doc_library.append({"name": source_name, "chunks": len(docs)})

                    scope = "all 8 architectures" if compare_mode else ARCH_INFO[selected_arch]["label"]
                    st.success(f"✅ {len(docs)} chunks → {scope}")

            except Exception as e:
                st.error(f"Ingestion failed: {e}")

    st.divider()

    # Query history
    if st.session_state.query_history:
        with st.expander(f"📋 History ({len(st.session_state.query_history)})", expanded=False):
            for item in reversed(st.session_state.query_history[-15:]):
                st.markdown(f"**Q:** {item['query'][:60]}{'…' if len(item['query']) > 60 else ''}")
                st.caption(f"{item['arch']} · {item['elapsed']:.2f}s")
                st.markdown("---")

    # Action buttons
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🗑️ Clear", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    with col2:
        if st.button("🔄 Reset", use_container_width=True):
            for key in ARCH_KEYS:
                sk = ARCH_INFO[key]["state_key"]
                st.session_state[sk].reset()
            st.session_state.ingested_archs = set()
            st.session_state.doc_library    = []
            st.success("Collections cleared.")
            st.rerun()
    with col3:
        if st.session_state.messages:
            md = export_chat_markdown()
            st.download_button(
                "💾 Export",
                data=md,
                file_name="rag_chat.md",
                mime="text/markdown",
                use_container_width=True,
            )

# ── Main ──────────────────────────────────────────────────────────────────────

st.title("Top 8 RAG Architectures in 2026")

# Architecture info card
info = ARCH_INFO[selected_arch]
with st.expander(f"{info['icon']} **{selected_arch}** — {info['tagline']}", expanded=False):
    col_a, col_b = st.columns([3, 1])
    with col_a:
        st.markdown("**How it works**")
        st.markdown(info["how"])
    with col_b:
        st.markdown("**Best for**")
        st.info(info["best_for"])

# Knowledge graph visualization (Graph RAG only)
if selected_arch == "02 Graph RAG (Knowledge Graphs)" and selected_arch in st.session_state.ingested_archs:
    graph_pipeline = st.session_state.graph_pipeline
    node_count = graph_pipeline.graph.number_of_nodes()
    edge_count = graph_pipeline.graph.number_of_edges()
    with st.expander(f"🕸️ Knowledge Graph — {node_count} entities · {edge_count} relationships", expanded=True):
        if node_count == 0:
            st.info("No entities extracted yet. Try ingesting a document with rich named entities.")
        else:
            html = graph_pipeline.render_graph_html()
            if html:
                components.html(html, height=540, scrolling=False)
                st.caption("Drag nodes to rearrange · Scroll to zoom · Hover for relationship labels")

st.divider()

# Chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("sources"):
            with st.expander(f"📎 {len(msg['sources'])} source chunks retrieved", expanded=False):
                for i, src in enumerate(msg["sources"]):
                    fname = os.path.basename(str(src.get("source", "Unknown")))
                    score_txt = f" · score {src['score']}" if src.get("score") is not None else ""
                    st.markdown(f"**Chunk {i + 1}** — `{fname}`{score_txt}")
                    st.caption(src["text"][:400] + "…" if len(src["text"]) > 400 else src["text"])
                    if i < len(msg["sources"]) - 1:
                        st.divider()
        if msg.get("eval"):
            ev = msg["eval"]
            ec1, ec2, ec3 = st.columns(3)
            ec1.metric(f"{score_color(ev['faithfulness'])} Faithfulness",    f"{ev['faithfulness']}/10")
            ec2.metric(f"{score_color(ev['relevance'])} Relevance",          f"{ev['relevance']}/10")
            ec3.metric(f"{score_color(ev['context_precision'])} Ctx Precision", f"{ev['context_precision']}/10")

# Chat input
if prompt := st.chat_input("Ask a question about your document…"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user", avatar="🧑"):
        st.markdown(prompt)

    with st.chat_message("assistant", avatar="🤖"):
        if compare_mode:
            st.markdown("### 🔍 Comparing All 8 Architectures")
            timings = {}
            answers = {}

            for arch_key in ARCH_KEYS:
                icon_a = ARCH_INFO[arch_key]["icon"]
                sk     = ARCH_INFO[arch_key]["state_key"]
                with st.status(f"{icon_a} {arch_key}…", expanded=False) as s:
                    t0 = time.time()
                    try:
                        answers[arch_key] = st.session_state[sk].query(prompt)
                    except Exception as e:
                        answers[arch_key] = f"⚠️ Error: {e}"
                    elapsed_i = time.time() - t0
                    timings[arch_key] = elapsed_i
                    s.update(label=f"{icon_a} {arch_key} — {elapsed_i:.2f}s", state="complete")

            timing_rows = "| Architecture | Time |\n|---|---|\n"
            for k in ARCH_KEYS:
                timing_rows += f"| {ARCH_INFO[k]['icon']} {k} | `{timings[k]:.2f}s` |\n"
            st.markdown(timing_rows)
            st.divider()

            for arch_key in ARCH_KEYS:
                icon_a = ARCH_INFO[arch_key]["icon"]
                with st.expander(f"{icon_a} **{arch_key}** — `{timings[arch_key]:.2f}s`", expanded=False):
                    st.markdown(answers[arch_key])

            response = "Comparison complete — expand each tab above to read the answers."
            st.session_state.messages.append({"role": "assistant", "content": response})

        else:
            collected_tokens  = []
            collected_sources = []
            answer  = ""
            elapsed = 0.0

            with st.status(f"🧠 {info['icon']} {selected_arch}…", expanded=True) as status:
                def on_event(event):
                    kind, content = event
                    if kind == "step":
                        st.write(f"✅ {content}")
                    elif kind == "token":
                        collected_tokens.append(content)
                    elif kind == "sources":
                        collected_sources.extend(content)

                t0 = time.time()
                try:
                    sk       = info["state_key"]
                    pipeline = st.session_state[sk]
                    answer   = pipeline.query(prompt, on_step=on_event)
                    elapsed  = time.time() - t0
                    status.update(label=f"✅ Done in {elapsed:.2f}s", state="complete", expanded=False)
                except Exception as e:
                    elapsed = time.time() - t0
                    answer  = f"⚠️ Pipeline error: {e}"
                    status.update(label="❌ Error", state="error", expanded=False)
                    st.error(answer)

            final_text = "".join(collected_tokens) if collected_tokens else answer
            st.caption(f"{info['icon']} **{info['label']}** · `{elapsed:.2f}s`")
            st.markdown(final_text)

            # Source citations
            if collected_sources:
                with st.expander(f"📎 {len(collected_sources)} source chunks retrieved", expanded=False):
                    for i, src in enumerate(collected_sources):
                        fname     = os.path.basename(str(src.get("source", "Unknown")))
                        score_txt = f" · score {src['score']}" if src.get("score") is not None else ""
                        st.markdown(f"**Chunk {i + 1}** — `{fname}`{score_txt}")
                        st.caption(src["text"][:400] + "…" if len(src["text"]) > 400 else src["text"])
                        if i < len(collected_sources) - 1:
                            st.divider()

            # RAG evaluation scorecard
            eval_scores = None
            if enable_eval and final_text and not final_text.startswith("⚠️"):
                with st.spinner("🧪 Evaluating answer quality…"):
                    eval_scores = evaluate_answer(prompt, collected_sources, final_text)
                ev = eval_scores
                ec1, ec2, ec3 = st.columns(3)
                ec1.metric(f"{score_color(ev['faithfulness'])} Faithfulness",       f"{ev['faithfulness']}/10")
                ec2.metric(f"{score_color(ev['relevance'])} Relevance",             f"{ev['relevance']}/10")
                ec3.metric(f"{score_color(ev['context_precision'])} Ctx Precision", f"{ev['context_precision']}/10")

            # Log to query history
            st.session_state.query_history.append({
                "query":   prompt,
                "arch":    selected_arch,
                "elapsed": elapsed,
                "answer":  final_text[:120],
            })

            # Persist message with sources + eval for replay in chat history
            st.session_state.messages.append({
                "role":    "assistant",
                "content": final_text,
                "sources": collected_sources,
                "eval":    eval_scores,
            })
